import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from datetime import datetime

def create_comprehensive_report():
    doc = Document()
    
    # Define project colors
    NAVY_BLUE = RGBColor(30, 58, 138)
    LIME_GREEN = RGBColor(132, 204, 22)
    DARK_GRAY = RGBColor(55, 65, 81)
    
    # --- 1. COVER PAGE ---
    logo_path = r'C:\Users\kingm\.gemini\antigravity\brain\b49f32a3-8278-4a8e-b029-91ca4129a713\athletepro_ai_logo_1778150332148.png'
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Inches(2.5)) # Slightly smaller logo
    
    doc.add_paragraph('\n') # Reduced space
    
    title = doc.add_paragraph('AthletePro AI')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(42) # Slightly smaller
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY_BLUE
    
    subtitle = doc.add_paragraph('RAG Tabanlı Kişiselleştirilmiş Sporcu Beslenme Platformu\nTeknik Mimari ve Uygulama Raporu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.runs[0]
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = LIME_GREEN
    
    doc.add_paragraph('\n' * 2) # Reduced space
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'Tarih: {datetime.now().strftime("%d %B %Y")}\nVersiyon: 1.0.6\nHazırlayan: AthletePro Teknik Ekibi\n')
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = DARK_GRAY
    
    doc.add_page_break()
    
    # --- 2. TABLE OF CONTENTS ---
    h_toc = doc.add_heading('İçindekiler', level=1)
    h_toc.runs[0].font.color.rgb = NAVY_BLUE
    
    toc_items = [
        "1. Proje Özeti ve Vizyon",
        "2. Sistem Mimarisi: RAG (Retrieval-Augmented Generation)",
        "3. Veri Mühendisliği ve Pipeline Süreçleri",
        "4. Yapay Zeka Modelleri ve Entegrasyonlar",
        "5. Kullanıcı Deneyimi (UX) ve Arayüz Tasarımı",
        "6. Örnek Veri Seti Analizi (Top 10 Tarif)",
        "7. Gelecek Yol Haritası ve Sonuç"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(0) # Remove space between TOC items
    
    # --- 3. PROJECT OVERVIEW ---
    h1 = doc.add_heading('1. Proje Özeti ve Vizyon', level=1)
    h1.runs[0].font.color.rgb = NAVY_BLUE
    doc.add_paragraph(
        "AthletePro AI, modern sporcunun en büyük ihtiyaçlarından biri olan 'veriye dayalı beslenme' sorununu çözmek amacıyla geliştirilmiş "
        "uçtan uca bir yapay zeka platformudur. Vizyonumuz, sadece bir tarif rehberi olmak değil, her sporcunun cebinde taşıyabileceği "
        "profesyonel bir beslenme koçu ve strateji merkezi haline gelmektir."
    )
    doc.add_paragraph(
        "Projenin temel amacı, karmaşık beslenme verilerini anlamlı ve uygulanabilir öğün planlarına dönüştürmektir. "
        "Yapay zeka desteğiyle, sporcuların antrenman yoğunluklarına ve kişisel hedeflerine en uygun içeriği saniyeler içinde "
        "erişilebilir kılıyoruz. Gelecekte, giyilebilir teknoloji entegrasyonları ile biyometrik verileri de işleyerek "
        "tamamen dinamik bir beslenme ekosistemi oluşturmayı hedefliyoruz."
    )
    
    # --- 4. SYSTEM ARCHITECTURE ---
    h2 = doc.add_heading('2. Sistem Mimarisi: RAG (Retrieval-Augmented Generation)', level=1)
    h2.runs[0].font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph(
        "RAG (Retrieval-Augmented Generation), Türkçesiyle 'Geri Getirme ile Güçlendirilmiş Üretim', "
        "büyük dil modellerinin (LLM) eğitim verilerinde bulunmayan güncel veya özel verilere erişmesini sağlayan "
        "hibrit bir mimaridir. Bu sistem, modelin halüsinasyon (yanlış bilgi üretme) riskini minimize ederken, "
        "yanıtların doğruluğunu ve güvenilirliğini en üst düzeye çıkarır."
    )
    
    # Detailed Steps
    rag_steps = [
        ("1. Veri Hazırlama ve İndeksleme (Indexing)", 
         "Ham veri (tarifler) küçük parçalara (chunks) ayrılır. Her parça 'all-MiniLM-L6-v2' modeli ile sayısal vektörlere "
         "dönüştürülür ve ChromaDB gibi bir vektör veritabanında saklanır."),
        ("2. Bilgi Getirme (Retrieval)", 
         "Kullanıcı bir soru sorduğunda, bu soru da vektöre dönüştürülür ve veritabanında 'vektör benzerlik araması' yapılarak "
         "en alakalı bilgi parçaları anlık olarak çekilir."),
        ("3. Bağlam Zenginleştirme (Augmentation)", 
         "Elde edilen bu özel bilgi parçaları, kullanıcının orijinal sorusu ile birleştirilerek LLM'e (GPT-4o) gönderilen "
         "prompt (istem) içine dahil edilir. Böylece model 'bilmediği' konularda elindeki bu bağlama dayanarak yanıt verir."),
        ("4. Yanıt Üretme (Generation)", 
         "LLM, kendisine sunulan bu özel bağlamı analiz ederek, sadece bu verilere sadık kalarak profesyonel ve "
         "doğru yanıtı üretir.")
    ]
    
    for step_title, step_desc in rag_steps:
        p = doc.add_paragraph()
        p.add_run(step_title).bold = True
        doc.add_paragraph(step_desc)
    
    # --- 5. DATA PIPELINE ---
    h3 = doc.add_heading('3. Veri Mühendisliği ve Pipeline Süreçleri', level=1)
    h3.runs[0].font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph(
        "Veri boru hattı (Data Pipeline), ham verinin akıllı bir bilgi kaynağına dönüştürülme sürecidir. Bu süreç şu aşamalardan oluşur:"
    )
    p_pipe = doc.add_paragraph()
    p_pipe.add_run("• Veri Toplama ve Temizleme: ").bold = True
    p_pipe.add_run("Özel seçilmiş 100 sporcu tarifi, metin formatına getirilerek gereksiz karakterlerden arındırılır.\n")
    p_pipe.add_run("• Parçalara Ayırma (Chunking): ").bold = True
    p_pipe.add_run("Anlamsal bütünlüğü korumak adına veriler 400 karakterlik bloklara bölünür.\n")
    p_pipe.add_run("• Vektörizasyon: ").bold = True
    p_pipe.add_run("Her bir blok, yüksek boyutlu vektörlere dönüştürülerek makine tarafından 'anlaşılabilir' hale getirilir.\n")
    p_pipe.add_run("• Kalıcı Depolama: ").bold = True
    p_pipe.add_run("Oluşturulan vektörler ChromaDB veritabanında saklanarak hızlı erişim için indekslenir.")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametre'
    hdr_cells[1].text = 'Değer / Strateji'
    data_rows = [
        ('Veri Kaynağı', '100 Adet Özel Seçilmiş Sporcu Tarifi'),
        ('Chunk Size', '400 Karakter'),
        ('Chunk Overlap', '50 Karakter'),
        ('Vector DB', 'ChromaDB (Local Persistence)'),
        ('Embedding Model', 'all-MiniLM-L6-v2 (384 Dim)')
    ]
    for param, val in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = val
    
    # --- 6. AI MODELS ---
    h4 = doc.add_heading('4. Yapay Zeka Modelleri ve Entegrasyonlar', level=1)
    h4.runs[0].font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph(
        "Sistemimiz, sektörün en gelişmiş yapay zeka modellerini bir orkestrasyon içinde kullanarak hibrit bir çözüm sunar:"
    )
    
    p_ai = doc.add_paragraph()
    p_ai.add_run("• GPT-4o (Reasoning & Text): ").bold = True
    p_ai.add_run(
        "Sistemin beynidir. RAG mimarisi ile gelen verileri analiz eder, sporcu beslenmesi kurallarına göre yorumlar "
        "ve kullanıcıya doğal dilde profesyonel yanıtlar üretir.\n"
    )
    p_ai.add_run("• DALL-E 3 (Visual Generation): ").bold = True
    p_ai.add_run(
        "Metin tabanlı tarifleri görselleştirir. Kullanıcının iştahını artıracak ve sunum kalitesini gösterecek "
        "yüksek çözünürlüklü, fotogerçekçi yemek fotoğrafları üretir.\n"
    )
    p_ai.add_run("• OpenAI API Entegrasyonu: ").bold = True
    p_ai.add_run(
        "Düşük gecikme süresi (low-latency) ve yüksek güvenlik standartları ile modellerin uygulama ile kesintisiz "
        "iletişimi sağlanır."
    )
    
    doc.add_page_break()
    
    # --- 7. UI/UX SECTION (WITH IMAGES) ---
    h5 = doc.add_heading('5. Kullanıcı Deneyimi (UX) ve Arayüz Tasarımı', level=1)
    h5.runs[0].font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph(
        "AthletePro AI, Streamlit kütüphanesi kullanılarak geliştirilmiş modern bir web arayüzüne sahiptir."
    )
    
    ui_images = [
        ('Dashboard', 'images/ui_screenshots/ui_dashboard.png'),
        ('AI Şef', 'images/ui_screenshots/ui_ai_chef.png'),
        ('Tarif Arşivi', 'images/ui_screenshots/ui_recipe_archive.png'),
        ('Alışveriş Listesi', 'images/ui_screenshots/ui_shopping_list.png'),
        ('Ayarlar', 'images/ui_screenshots/ui_settings.png')
    ]
    
    for title, path in ui_images:
        # Smaller headings for UI images and keep them close
        rt = doc.add_paragraph(title)
        rt.runs[0].font.bold = True
        rt.runs[0].font.size = Pt(12)
        rt.runs[0].font.color.rgb = LIME_GREEN
        rt.paragraph_format.space_after = Pt(2)
        
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4.5)) # Smaller images to fit more
            doc.add_paragraph().paragraph_format.space_after = Pt(5) # Minimal space
            
    doc.add_page_break()
    
    # --- 8. SAMPLE DATA (10 RECIPES) ---
    h6 = doc.add_heading('6. Örnek Veri Seti Analizi (Top 10 Tarif)', level=1)
    h6.runs[0].font.color.rgb = NAVY_BLUE
    
    md_recipes_path = r'c:\Users\kingm\.gemini\antigravity\scratch\ai_recipe_generator\data\saglikli_tarifler.txt'
    if os.path.exists(md_recipes_path):
        with open(md_recipes_path, 'r', encoding='utf-8') as f:
            content = f.read()
            recipes = content.split('\n\n')[:10]
            
            for i, recipe in enumerate(recipes):
                if not recipe.strip(): continue
                lines = recipe.strip().split('\n')
                rt = doc.add_heading(f"{i+1}. {lines[0].replace('Tarif ', '')}", level=2)
                rt.runs[0].font.color.rgb = LIME_GREEN
                rt.paragraph_format.space_before = Pt(5)
                for r_line in lines[1:]:
                    p = doc.add_paragraph(r_line)
                    p.paragraph_format.space_after = Pt(2) # Compact lines
                
    # --- 9. ROADMAP & CONCLUSION ---
    h7 = doc.add_heading('7. Gelecek Yol Haritası ve Sonuç', level=1)
    h7.runs[0].font.color.rgb = NAVY_BLUE
    doc.add_paragraph("Planlanan iyileştirmeler: Kullanıcı Profilleme, Market Entegrasyonu ve Vision desteği.")
    
    # --- FOOTER ---
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "AthletePro AI | Teknik Mimari Dokümantasyonu"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_name = 'AthletePro_Teknik_Rapor_Compact.docx'
    doc.save(output_name)
    print(f"Kompakt rapor başarıyla oluşturuldu: {output_name}")

if __name__ == "__main__":
    create_comprehensive_report()
