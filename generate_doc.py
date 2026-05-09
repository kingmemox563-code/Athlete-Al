from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('AthletePro AI: Teknik Sunum Dosyası', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    p = doc.add_paragraph('Bu dosya, AthletePro AI projesinin teknik altyapısını, kullanılan teknolojileri ve sistem mimarisini detaylandırmaktadır.')
    p.italic = True
    
    # Sections
    sections = [
        ("1. Sistem Mimarisi: RAG (Retrieval-Augmented Generation)", 
         "Bu uygulama klasik bir chatbot değildir; bir RAG mimarisi üzerine kurulmuştur. Yapay zeka cevap vermeden önce onaylı 100 tariflik özel veritabanını tarar, en uygun bilgiyi geri getirme adımıyla alır ve cevabı bu gerçek verilere dayanarak oluşturur."),
        
        ("2. Veri Kaynağı ve İşleme (Data Pipeline)", 
         "- Veri Seti: 100 adet sağlıklı, yüksek proteinli ve sporcu odaklı tariften oluşan özel bir korpus (saglikli_tarifler.txt).\n"
         "- Temizleme: Ham metin verileri; başlık, malzemeler, makro değerler ve hazırlanış adımları şeklinde normalize edildi.\n"
         "- Parçalama: RecursiveCharacterTextSplitter kullanılarak veriler küçük parçalara (chunks) ayrıldı."),
        
        ("3. Veri Tabanı ve Vektör Arama", 
         "- Vektör Veri Tabanı: ChromaDB. Veriler matematiksel vektörler olarak saklanır.\n"
         "- Embeddings: sentence-transformers/all-MiniLM-L6-v2 (HuggingFace) modeli kullanıldı. Bu model, kelimelerin anlamlarını 384 boyutlu bir uzayda temsil eder."),
        
        ("4. Yapay Zeka Modelleri (LLM & Vision)", 
         "- Ana Beyin: OpenAI GPT-4o. Gelen verileri analiz edip kişiselleştirilmiş menüleri üretir.\n"
         "- Görsel Üretim: OpenAI DALL-E 3. Tarif içeriğine göre dinamik ve yüksek çözünürlüklü yemek fotoğrafları üretir.\n"
         "- Strateji: DALL-E komutları tarif malzemelerine göre otomatik olarak optimize edilir."),
        
        ("5. Yazılım Teknolojileri ve Kütüphaneler", 
         "- Frontend/UI: Streamlit (Modern glassmorphism tasarım).\n"
         "- Orchestration: LangChain (AI ve veritabanı akış yönetimi).\n"
         "- PDF Motoru: ReportLab (Platypus) - Unicode ve görsel destekli.\n"
         "- Veri Yönetimi: Pydantic ile yapılandırılmış JSON çıktısı."),
        
        ("6. Öne Çıkan Özellikler", 
         "1. Kişiselleştirilmiş Makro Hesaplama: BMR ve TDEE bazlı protein hedefleri.\n"
         "2. Kalıcı Görsel Hafıza: images/ klasöründe saklanan görseller.\n"
         "3. Akıllı Alışveriş Listesi: Tek tıkla malzeme listesi oluşturma.\n"
         "4. Profesyonel Raporlama: Resimli ve makro tablolu PDF çıktısı.")
    ]
    
    for head, body in sections:
        h = doc.add_heading(head, level=1)
        h_run = h.runs[0]
        h_run.font.color.rgb = RGBColor(37, 99, 235)
        
        doc.add_paragraph(body)
    
    # Summary
    doc.add_heading('Sunum Özeti', level=2)
    summary = doc.add_paragraph("Bu proje; ChromaDB vektör veritabanı ve LangChain framework'ü kullanarak, sporcuların ihtiyaç duyduğu bilimsel beslenme verilerini GPT-4o'nun üretim gücüyle birleştiren, DALL-E 3 ile görselleştirilmiş ve ReportLab ile profesyonel raporlama sunan uçtan uca bir SaaS platformu prototipidir.")
    summary.style = 'Quote'
    
    doc.save('AthletePro_Teknik_Sunum.docx')
    print("Belge başarıyla oluşturuldu: AthletePro_Teknik_Sunum.docx")

if __name__ == "__main__":
    create_tech_doc()
